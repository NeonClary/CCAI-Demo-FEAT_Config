# NEON AI (TM) SOFTWARE, Software Development Kit & Application Framework
# All Rights Reserved 2008-2025
# Licensed under the BSD 3-Clause License
# https://opensource.org/licenses/BSD-3-Clause
#
# Copyright (c) 2008-2025, Neongecko.com Inc.
#
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
# 1. Redistributions of source code must retain the above copyright notice,
#    this list of conditions and the following disclaimer.
# 2. Redistributions in binary form must reproduce the above copyright notice,
#    this list of conditions and the following disclaimer in the documentation
#    and/or other materials provided with the distribution.
# 3. Neither the name of the copyright holder nor the names of its contributors
#    may be used to endorse or promote products derived from this software
#    without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE
# ARE DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE
# LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR
# CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF
# SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS
# INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN
# CONTRACT, STRICT LIABILITY, OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE)
# ARISING IN ANY WAY OUT OF THE USE OF THIS SOFTWARE, EVEN IF ADVISED OF THE
# POSSIBILITY OF SUCH DAMAGE.

import re
from io import BytesIO
from typing import Any, Dict, List, Tuple, Union

from docx import Document
from fastapi.responses import StreamingResponse
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


def format_messages_for_export(messages: List[Dict[str, Any]]) -> str:
    """
    Convert chat messages into a structured exportable string.
    """
    return "\n\n".join([f"{m['role']}:\n{m['content'].strip()}" for m in messages])


def generate_txt_file(text: str) -> BytesIO:
    buffer = BytesIO()
    buffer.write(text.encode("utf-8"))
    buffer.seek(0)
    return buffer


def generate_docx_file(text: str) -> BytesIO:
    doc = Document()
    for block in text.split("\n\n"):
        doc.add_paragraph(block)
        doc.add_paragraph("")  # spacing

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _clean_text_for_pdf(text: str) -> str:
    """
    Clean text for PDF generation to handle special characters and formatting.
    """
    text = text.replace('\u2019', "'")
    text = text.replace('\u2018', "'")
    text = text.replace('\u201c', '"')
    text = text.replace('\u201d', '"')
    text = text.replace('\u2013', '-')
    text = text.replace('\u2014', '-')

    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)

    return text


def generate_pdf_file(text: str) -> BytesIO:
    """
    Improved PDF generation with proper text wrapping and formatting.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )

    styles = getSampleStyleSheet()

    role_style = ParagraphStyle(
        name="RoleStyle",
        parent=styles["Normal"],
        fontSize=12,
        fontName="Helvetica-Bold",
        spaceAfter=6,
        textColor='blue'
    )

    content_style = ParagraphStyle(
        name="ContentStyle",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica",
        leading=14,
        spaceAfter=12,
        leftIndent=20
    )

    story: List[Any] = []

    blocks = text.split("\n\n")

    for block in blocks:
        if not block.strip():
            continue

        clean_block = _clean_text_for_pdf(block.strip())

        lines = clean_block.split('\n', 1)
        if len(lines) > 1 and lines[0].strip().endswith(':'):
            role = lines[0].strip()
            content = lines[1].strip() if len(lines) > 1 else ""

            story.append(Paragraph(role, role_style))

            if content:
                content_paragraphs = content.split('\n')
                for para in content_paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), content_style))
        else:
            story.append(Paragraph(clean_block, content_style))

        story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _render_rich_text(text: str) -> str:
    """
    Convert markdown-style **bold** to reportlab <b>bold</b> tags.
    """
    return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)


def generate_pdf_file_from_blocks(blocks: List[Dict[str, Any]]) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)

    styles = getSampleStyleSheet()

    heading_style = ParagraphStyle(
        name="Heading",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        spaceAfter=20
    )

    paragraph_style = ParagraphStyle(
        name="Paragraph",
        parent=styles["Normal"],
        alignment=TA_LEFT,
        leading=14,
        spaceAfter=8
    )

    list_style = ParagraphStyle(
        name="List",
        parent=styles["Normal"],
        leftIndent=20,
        leading=14
    )

    story: List[Any] = []

    for block in blocks:
        if block["type"] == "heading":
            story.append(Paragraph(_render_rich_text(block["text"]), heading_style))
            story.append(Spacer(1, 12))

        elif block["type"] == "paragraph":
            story.append(Paragraph(_render_rich_text(block["text"]), paragraph_style))

        elif block["type"] == "list":
            list_items = [
                ListItem(Paragraph(_render_rich_text(item), list_style))
                for item in block["items"]
            ]
            story.append(
                ListFlowable(
                    list_items,
                    bulletType="1" if block.get("style") == "numbered" else "bullet",
                    start="1" if block.get("style") == "numbered" else None,
                    leftIndent=20
                )
            )
            story.append(Spacer(1, 8))

    doc.build(story)
    buffer.seek(0)
    return buffer


def export_chat_as_file(
    content: Union[str, List[Dict[str, Any]]], format: str
) -> Tuple[BytesIO, str, str]:
    """
    Export either a list of chat messages or a summary string to the specified format.
    """
    if isinstance(content, list):
        text = format_messages_for_export(content)
    elif isinstance(content, str):
        text = content.strip()
    else:
        raise ValueError("Unsupported content type")

    if format == "txt":
        return generate_txt_file(text), "chat_export.txt", "text/plain"

    elif format == "docx":
        return generate_docx_file(text), "chat_export.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif format == "pdf":
        return generate_pdf_file(text), "chat_export.pdf", "application/pdf"

    else:
        raise ValueError(f"Unsupported export format: {format}")


def prepare_export_response(
    content: Union[str, List[Dict[str, Any]]],
    format: str,
    filename_prefix: str = "chat_export"
) -> StreamingResponse:
    """
    Prepare a StreamingResponse for export, using the given filename prefix.
    """
    stream, filename, media_type = export_chat_as_file(content, format)

    final_filename = filename.replace("chat_export", filename_prefix)

    return StreamingResponse(
        stream,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={final_filename}"}
    )
